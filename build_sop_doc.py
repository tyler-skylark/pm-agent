#!/usr/bin/env python3
"""Render the Skylark AV SOP context as a clean, human-readable Word doc."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


SECTIONS = [
    ("Project Types", [
        "Skylark categorizes every active SKY- project into one of two types.",
        ("**Design Contract** — Pre-execution, design/engineering only. The project name "
         "includes \"(Design Contract)\". TBD fields are acceptable. Full SOP compliance "
         "(install scheduling, labor, procurement) is NOT expected at this stage."),
        ("**Standard Project** — Full execution project. All description fields are "
         "required, no TBDs allowed, full SOP applies."),
    ]),
    ("Project Description (Required Fields)", [
        "Every active SKY- project must have exactly these five fields in its description:",
        "• Client Contact",
        "• Job Location",
        "• Skylark PM",
        "• Engineer",
        "• On-Site Lead",
        "No dates belong in the description — dates live only in schedule-tagged todos.",
        "TBD is only acceptable on Design Contract projects, never on Standard Projects.",
    ]),
    ("Schedule Tag System", [
        ("Skylark uses bracketed tags in todo titles to indicate which phase a scheduled "
         "item belongs to. Every scheduled todo must carry one of these tags AND a due date."),
        "Valid tags: [PM-SCHED], [ENG-SCHED], [PROC-SCHED], [SHOP-SCHED], [LOG-SCHED], [ONS-SCHED], [COM-SCHED], [FUT-SCHED]",
        "RULE: Any incomplete schedule-tagged todo without a due date is a flag.",
    ]),
    ("Engineering Milestone System", [
        ("Engineering runs two coordinated tracks — Construction DD/CD and AV Systems. "
         "These tracks are related but distinct. Procurement releases happen in stages "
         "as engineering confidence increases. Once a package reaches IFC it is in "
         "revision-only mode; all changes tracked via Revision IDs and revision clouds."),
        "Key terms:",
        "• DD — Design Development issue phase for construction coordination",
        "• CD — Post-DD coordination phase driven by trades, client changes, or field conditions",
        "• IFC — Issued for Construction (final released package)",
        "• Issue Set — A formally published milestone package",
        "• Revision — A tracked change to an issued package",
    ]),
    ("Construction DD/CD Track [ENG-SCHED]", [
        ("**Milestone 1** — Early GC coordination: weight loads, heat loads, AV spaces "
         "identified, total power requirements, \"provided by others\" callouts, rough "
         "coordination mockups. Deliverable: DD/CD Issue Set."),
        ("**Milestone 2** — Further GC coordination: circuit J-box locations, furniture "
         "mockups. Deliverable: DD/CD Issue Set."),
        ("**Milestone 3** — Construction DD effectively complete. Must include a complete "
         "and coherent coordination package, ready for internal review then immediate "
         "client review."),
        "    Review sequence: (1) Internal design review → (2) Client review",
        ("**Milestone 4** — 100% GC handoff to trades. Baseline issue for construction "
         "coordination, design development complete. All future changes must come from "
         "coordination feedback, field conditions, or client-driven changes — not "
         "unfinished engineering work. Deliverable: 100% DD/CD Construction Issue Set."),
        ("**CD Phase** — Responds to changes after DD baseline. CD changes are reactions "
         "to other trades, not open-ended design. The same 100% DD package may be re-issued "
         "to meet GC/architect milestones; those external milestones do not create new "
         "internal engineering phases. Revisions must be tracked clearly."),
    ]),
    ("AV Systems Track [ENG-SCHED]", [
        ("**25% AV Systems** — Establish big-picture system layout and surface risk early. "
         "Includes file structure, equipment blocks layout, system roughly built by room, "
         "big-picture system view, early lines/cables, Order Detail items represented. "
         "Internal use: Concept Drawing Comparison Review, PM checkpoint, early discovery "
         "of missing gear/scope gaps. Procurement: Release 1 for major long-lead items."),
        ("**50% AV Systems** — Package far enough for engineering coordination and shop "
         "planning. Includes wires drawn, rack elevations, power plan, Custom Panels, "
         "overall system layout complete. Deliverables: first-pass Design Review, "
         "package may be given to shop for planning, Custom Panels client signoff. "
         "Procurement: Release 2 for additional equipment."),
        ("**75% AV Systems** — Essentially complete systems design. Includes systems "
         "design complete, wire numbers complete, system logic complete, package ready "
         "for systems review. Review: internal systems design review. Procurement: "
         "Release 3 for remaining equipment."),
        ("**100% AV Systems** — Complete physical implementation package. Includes "
         "physical design complete, speaker locations, LED wall details, TV placement, "
         "rigging drawings, cable pull schedules. Review: internal review with install "
         "team involvement, client review as needed. Procurement: Release 4 for rigging "
         "and physical infrastructure."),
        ("**IFC (Issued for Construction)** — Final released package for execution. "
         "Revision-only mode from this point forward."),
    ]),
    ("Engineering SOP Violations to Flag", [
        "• Engineering doing new or open-ended design work after Milestone 4 / CD phase (should be reactions only)",
        "• No review sequence at Milestone 3 (internal → client missing)",
        "• Procurement released before its corresponding AV Systems milestone",
        "• IFC package being edited without tracked Revision IDs",
        "• Missing milestone deliverables (no Issue Set published at M1 / M2 / M4)",
    ]),
    ("Closeout Schedule Anchors", [
        "• Punch List Walkthrough [PM-SCHED] — 48 hours before end of install",
        "• Client Sign-Off [PM-SCHED] — before pulling off the job",
        "• As-Built Package [ENG-SCHED] — 2 weeks after open",
        "• Post-Mortem [PM-SCHED] — 1 week after open",
        "• Project Closed [PM-SCHED] — 90 days after open",
    ]),
    ("Labor Scheduling", [
        ("[LABOR] todos live in the Labor Scheduling todoset. The format is: "
         "Name | Role | Status [LABOR]."),
        "The description must include Flights, Hotel, Per-Diem, and Car Rental.",
        "Any upcoming trip with missing travel info is a flag.",
        ("EVERY confirmed [ONS-SCHED] onsite trip must have corresponding [LABOR] "
         "todos documenting who is going onsite for that trip. This is non-negotiable — "
         "onsite work without documented labor is a SOP violation."),
        ("Tie labor to the trip by date proximity: the [LABOR] dates should overlap "
         "or align with the [ONS-SCHED] trip date. For each [ONS-SCHED] trip with a "
         "real due date, there must be at least one [LABOR] todo in the Labor "
         "Scheduling todoset that covers that trip's date range."),
        "Flag if:",
        "• A confirmed [ONS-SCHED] trip has zero [LABOR] todos covering it",
        ("• A [LABOR] todo is missing Flights / Hotel / Per-Diem / Car Rental when "
         "the trip is within 14 days"),
        "• A [LABOR] todo exists but has no due date or assignee",
    ]),
    ("Required Client-Visible Todo Lists (Standard Projects only)", [
        "These three todo lists must be set to \"The client sees this\" on every Standard Project:",
        "• Onsite Phase",
        "• Commissioning Phase",
        "• Closeout Phase",
        ("Any list missing from the project, or present but not client-visible, is "
         "flagged as a Client Visibility SOP violation."),
    ]),
    ("Pre-Mobilization Gate", [
        "GO/NO-GO check is required 14 days AND 7 days before mobilization.",
        "If an [ONS-SCHED] item is due in <14 days with no GO/NO-GO evidence, flag.",
    ]),
    ("Required Logistics Todos (Standard Projects with confirmed onsite date)", [
        ("A project can have multiple onsite trips. Each confirmed [ONS-SCHED] todo "
         "(with a real due date, not TBD) represents one trip and requires its own "
         "complete set of three logistics todos:"),
        "1. \"Equipment/Materials Arrive Onsite [LOG-SCHED]\"",
        "2. \"Verify Equipment and Materials [LOG-SCHED]\"",
        "3. \"Return Trip to Home Shop (Excess Equipment/Materials) [LOG-SCHED]\"",
        "Example: 3 onsite trips → 3 of each logistics todo, all with due dates.",
        "Flag as SOP Deviation if:",
        "• Any logistics todo count is less than the number of confirmed [ONS-SCHED] todos",
        "• Any logistics todo is present but missing its due date",
        "This check applies only to Standard Projects, not Design Contracts.",
    ]),
    ("Onsite \"Attention Needed\" List (Field RFIs)", [
        ("Inside the Onsite Tasks todoset there is a todolist called \"Attention "
         "Needed.\" These are field RFIs raised by the install team that need fast "
         "resolution."),
        ("RULE: Any open todo in this list whose updated_at timestamp is more than "
         "24 hours old (no comment, status change, or assignment update in 24h) is "
         "flagged as a stale field RFI."),
        ("Use updated_at (last touched), not created_at. A field RFI created 3 days "
         "ago but updated this morning is fine; one created today but untouched for "
         "25 hours is not."),
        "Stale RFI flags should include the title, project, hours since last activity, assignee (no assignee is a worse flag), and a direct link.",
    ]),
    ("Onsite Trip Companion Items (Standard Projects)", [
        "Every confirmed [ONS-SCHED] installation trip should be followed by these companion items, all on the project schedule:",
        "• Commissioning task (typically [COM-SCHED]) — system commissioning after install",
        "• Client Training — as needed for the scope (cameras, audio, lighting, control surfaces, etc.)",
        "• Punch List Walkthrough [PM-SCHED] — 48 hours before end of install",
        "Flag any [ONS-SCHED] install trip that is missing a following commissioning todo, training (where scope clearly requires it), or a punch list walkthrough.",
    ]),
    ("Punch List Phase Detection & Client Visibility", [
        ("A project enters Punch List Phase when there is an [ONS-SCHED] todo whose "
         "title contains \"Installation (Punch List)\" or \"Punch List\" "
         "(e.g. Installation (Punch List) [ONS-SCHED])."),
        "Once in Punch List Phase:",
        ("1. All remaining open punch list items in the Onsite Tasks, Engineering Tasks, "
         "and Commissioning Tasks todosets must be moved to the Closeout Tasks todoset "
         "(moved, not duplicated)."),
        ("2. The Closeout list containing these items must be set to client-visible "
         "(\"The client sees this\") so the client can track punch list progress."),
        ("3. Any item still sitting in Onsite, Engineering, or Commissioning during "
         "Punch List Phase is a SOP deviation. Each gets flagged with its title, "
         "current todoset, and the action needed (\"move to Closeout\")."),
        ("Punch-list-style items include touch-ups, corrections, deficiencies, and the "
         "list of items the client called out. These belong in Closeout once the "
         "project is in Punch List Phase."),
    ]),
    ("Communication Rules", [
        "• Client posts → \"Client Communication\" message board only",
        "• Internal updates → \"Internal Coordination\" message board only",
        "• Decisions and actions from calls must be logged in Basecamp",
    ]),
    ("Closeout", [
        ("A project is overdue for closure if \"Client First Open [PM-SCHED]\" "
         "passed more than 90 days ago AND \"Project Closed in Basecamp [PM-SCHED]\" "
         "is still incomplete."),
    ]),
    ("Google Drive Job Folder Compliance", [
        ("Every Standard Project should have a Google Drive project folder named "
         "SKY-XXXX … inside the client's folder under \"Skylark Jobs.\" The template "
         "calls for these top-level folders:"),
        "• Contract Docs",
        "• Engineering",
        "• Proposals",
        "• Vendor Docs",
        "Each top-level folder has its own required subfolders.",
        ("REAL-WORLD RULE: The template is a guide, not a contract. Before flagging "
         "something missing, look at the actual folder tree and ask, \"is there "
         "functional evidence this requirement is met?\""),
        "Examples:",
        ("• Template wants \"Signed Contracts/\" but the tree shows \"Contract "
         "Revisions/\" containing the executed contract PDF and a \"Signed Documents/\" "
         "subfolder — that IS the signed contract on file. Don't flag."),
        ("• Template wants \"Insurance Documents/\" but the tree shows \"Insurance "
         "Docs/\" with a COI PDF inside — same thing. Don't flag."),
        ("• Template wants \"Onsite Pictures/\" but the tree shows an empty \"Onsite "
         "Photos/\" — flag as empty if the project is past the onsite phase."),
        ("• Template wants \"Signed Contracts/\" and nothing in the tree looks like a "
         "contract document — flag it, this is a real gap."),
        ("Cross-reference against Basecamp phase. Missing a signed contract on a "
         "design-phase project is normal. Missing it on a project already in onsite "
         "phase is a red flag."),
    ]),
]


def add_styled_paragraph(doc, text, bold=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bold_inline(doc, text, size=11):
    """Render a paragraph that uses **bold** markdown markers as inline bold runs."""
    p = doc.add_paragraph()
    parts = text.split("**")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.size = Pt(size)
        if i % 2 == 1:
            run.bold = True
    p.paragraph_format.space_after = Pt(6)
    return p


def build_doc(out_path):
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Skylark AV")
    run.bold = True
    run.font.size = Pt(28)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Operations Standards & Project Management SOP")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Updated {date.today().strftime('%B %d, %Y')}")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    intro = doc.add_paragraph()
    run = intro.add_run(
        "This document defines the operating standards used across all Skylark "
        "AV projects in Basecamp and Google Drive. It is the source-of-truth "
        "the Rick Stamen PM agent uses to evaluate project health, and the "
        "playbook the project team is expected to follow. Sections cover "
        "project types, scheduling conventions, engineering milestones, "
        "logistics, communication rules, closeout, and file management."
    )
    run.font.size = Pt(11)
    intro.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Sections
    for heading, items in SECTIONS:
        h = doc.add_heading(heading, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        for line in items:
            if "**" in line:
                add_bold_inline(doc, line)
            else:
                add_styled_paragraph(doc, line)

    # Footer-style note
    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run(
        "Maintained alongside the Rick Stamen PM agent — "
        "any change to this document should be reflected in the agent's SOP context, and vice versa."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    out = Path.home() / "Desktop" / "Skylark_AV_Operations_SOP.docx"
    build_doc(out)
    print(f"Wrote: {out}")
